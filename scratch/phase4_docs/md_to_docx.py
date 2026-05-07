"""Convert a Phase 4 doc markdown file into a .docx file based on a template.

Strategy:
- Open the template .docx (inherits its styles: fonts, headings, body).
- Clear all body content.
- Walk the markdown line-by-line and append paragraphs / tables.

Markdown subset supported (matches what the Phase 4 docs actually use):
- Headings: # / ## / ### / ####
- Paragraphs (blank-line separated)
- GFM-style tables (| ... | with a |---| separator row)
- Fenced and inline code (rendered as monospace runs; no fences in our docs)
- Bullet lists (- item)
- Bold (**...**) and italic (*...*) inline
- Block quotes ignored (none in our docs)
- Horizontal rule (---) emitted as a section break paragraph
"""

from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*[:\-]+\s*(\|\s*[:\-]+\s*)+\|?\s*$")
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
HR_RE = re.compile(r"^\s*-{3,}\s*$")
IMAGE_RE = re.compile(r"^\s*!\[(.*)\]\(([^)]+)\)\s*$")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")


def clear_document_body(doc: Document) -> None:
    """Remove all paragraphs and tables from the document body, preserving styles."""
    body = doc.element.body
    # Remove every direct child except the section properties at the end.
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def parse_inline(text: str):
    """Yield (kind, content) tuples for an inline-formatted text segment.

    kind is one of: "plain", "bold", "italic", "code". Bold and italic content
    is returned as raw markdown so callers can recurse to handle nesting like
    bold-with-inline-code. One-pass: at each position scan for the next
    occurring marker and emit a plain segment up to it followed by the
    marker's match.
    """
    out = []
    i = 0
    n = len(text)
    while i < n:
        candidates = []  # (start, end, kind, content)
        m = INLINE_CODE_RE.search(text, i)
        if m:
            candidates.append((m.start(), m.end(), "code", m.group(1)))
        m = BOLD_RE.search(text, i)
        if m:
            candidates.append((m.start(), m.end(), "bold", m.group(1)))
        m = ITALIC_RE.search(text, i)
        if m:
            candidates.append((m.start(), m.end(), "italic", m.group(1)))

        if not candidates:
            out.append(("plain", text[i:]))
            break

        # Earliest by start; bold beats italic at the same start.
        candidates.sort(
            key=lambda c: (
                c[0],
                0 if c[2] == "bold" else (1 if c[2] == "code" else 2),
            )
        )
        start_, end_, kind, content = candidates[0]
        if start_ > i:
            out.append(("plain", text[i:start_]))
        out.append((kind, content))
        i = end_

    return out


def _emit_run(paragraph, content: str, *, bold: bool, italic: bool, code: bool) -> None:
    run = paragraph.add_run(content)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_inline_runs(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    """Render an inline-formatted string into the paragraph.

    Recurses into bold/italic spans so that nested inline-code spans
    (e.g. ``**Foo (`bar`)**``) emit a run that is both bold and code,
    rather than leaking the internal code-span sentinel into the output.
    """
    for kind, content in parse_inline(text):
        if kind == "bold":
            add_inline_runs(paragraph, content, bold=True, italic=italic)
        elif kind == "italic":
            add_inline_runs(paragraph, content, bold=bold, italic=True)
        elif kind == "code":
            _emit_run(paragraph, content, bold=bold, italic=italic, code=True)
        else:  # "plain"
            _emit_run(paragraph, content, bold=bold, italic=italic, code=False)


def add_heading(doc: Document, level: int, text: str) -> None:
    # python-docx Heading 1..4 maps fine to template's heading styles.
    p = doc.add_heading(level=level)
    add_inline_runs(p, text)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    add_inline_runs(p, text)


def add_bullet(doc: Document, text: str) -> None:
    # Fall back to a plain paragraph with a bullet glyph; not every template
    # defines a "List Bullet" style.
    try:
        p = doc.add_paragraph(style="List Bullet")
        add_inline_runs(p, text)
    except KeyError:
        p = doc.add_paragraph()
        p.add_run("• ")  # bullet + non-breaking space
        add_inline_runs(p, text)


def add_image(doc: Document, image_path: Path, caption: str | None) -> None:
    """Embed a centered image at native aspect, with an optional italic caption.

    Width is capped at 6.0 inches to fit the standard letter-page text column.
    The caption is rendered as a centered italic paragraph immediately below
    the image, in the form "Figure N. <caption>" if the caller-supplied
    caption already includes the figure number, or just "<caption>" otherwise
    — figure numbering is the responsibility of the markdown source.
    """
    if not image_path.exists():
        # Don't crash; embed a placeholder paragraph so the missing image is
        # visible in the output rather than silently dropped.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[missing image: {image_path}]")
        run.italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.0))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cp.add_run(caption)
        crun.italic = True


def add_table(doc: Document, rows: list[list[str]], align_chars: list[str]) -> None:
    """Add a table from parsed cells. align_chars is per-column ':' marker for centering / right-align."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""  # clear default empty paragraph contents
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text)
            if r_idx == 0:
                # Bold the header row.
                for run in p.runs:
                    run.bold = True
            # Apply alignment from separator row markers.
            ac = align_chars[c_idx] if c_idx < len(align_chars) else "l"
            if ac == "r":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif ac == "c":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_table_separator(line: str) -> list[str]:
    """Parse a markdown table separator row and return per-column alignment chars (l/c/r)."""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    out = []
    for cell in cells:
        if cell.startswith(":") and cell.endswith(":"):
            out.append("c")
        elif cell.endswith(":"):
            out.append("r")
        elif cell.startswith(":"):
            out.append("l")
        else:
            out.append("l")
    return out


def parse_table_row(line: str) -> list[str]:
    """Parse a markdown table data row into a list of cells."""
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [c.strip() for c in inner.split("|")]


def convert(md_path: Path, template_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document(str(template_path))
    clear_document_body(doc)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line: skip.
        if not stripped:
            i += 1
            continue

        # Horizontal rule.
        if HR_RE.match(stripped):
            # Render as an empty centered paragraph with three em-dashes.
            p = doc.add_paragraph("———")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Heading.
        m = HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            add_heading(doc, level, m.group(2).strip())
            i += 1
            continue

        # Image: ![caption](path). Path may be absolute or relative to the
        # markdown file's directory.
        m = IMAGE_RE.match(stripped)
        if m:
            caption = m.group(1).strip()
            raw_path = m.group(2).strip()
            img_path = Path(raw_path)
            if not img_path.is_absolute():
                img_path = (md_path.parent / raw_path).resolve()
            add_image(doc, img_path, caption or None)
            i += 1
            continue

        # Bullet list.
        if BULLET_RE.match(stripped):
            while i < n and BULLET_RE.match(lines[i].strip()):
                bm = BULLET_RE.match(lines[i].strip())
                add_bullet(doc, bm.group(1))
                i += 1
            continue

        # Table: a row line followed by a separator on the next line.
        if (
            TABLE_ROW_RE.match(line)
            and i + 1 < n
            and TABLE_SEP_RE.match(lines[i + 1])
        ):
            header = parse_table_row(lines[i])
            align = parse_table_separator(lines[i + 1])
            i += 2
            data_rows = []
            while i < n and TABLE_ROW_RE.match(lines[i]):
                data_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, [header] + data_rows, align)
            continue

        # Default: paragraph. Collect consecutive non-blank, non-special lines.
        para_lines = [line]
        i += 1
        while i < n:
            next_line = lines[i]
            ns = next_line.strip()
            if not ns:
                break
            if HEADING_RE.match(ns) or BULLET_RE.match(ns) or HR_RE.match(ns):
                break
            if IMAGE_RE.match(ns):
                break
            if TABLE_ROW_RE.match(next_line) and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1]):
                break
            para_lines.append(next_line)
            i += 1
        add_paragraph(doc, " ".join(l.strip() for l in para_lines))

    doc.save(str(out_path))
    print(f"wrote {out_path}")


def main(argv: list[str]) -> int:
    if len(argv) != 4:
        print(
            "usage: md_to_docx.py <md-source> <template-docx> <output-docx>",
            file=sys.stderr,
        )
        return 2
    md_path = Path(argv[1]).resolve()
    template_path = Path(argv[2]).resolve()
    out_path = Path(argv[3]).resolve()
    convert(md_path, template_path, out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
